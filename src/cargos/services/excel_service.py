"""
Service classes for Excel processing and file generation.
"""
import pandas as pd
import logging
from pathlib import Path
from typing import List, Optional, Dict, Any
from cargos.core.models import ExcelData, GenerationResult, AppConfig, ExcelValidationResult, WorksheetMetadata, WorksheetParsingResult, GenerationOptions
from cargos.core.validators import TemplateValidator
from cargos.services.unified_config_service import UnifiedConfigService
from cargos.core.constants import (
    METADATA_ROW_FECHA_SOLICITUD, METADATA_COL_FECHA_SOLICITUD,
    METADATA_ROW_TIENDA, METADATA_COL_TIENDA,
    METADATA_ROW_ADMINISTRADOR, METADATA_COL_ADMINISTRADOR,
    LOCATION_ROW, OCCUPATION_ROW,
    HEADER_ROW, DATA_START_ROW, IGNORE_COLUMN_INDEX, MAIN_DATA_END_COLUMN,
    UNIFORM_DATA_START_ROW, UNIFORM_DATA_START_COLUMN, UNIFORM_DATA_END_COLUMN,
    UNIFORM_COLUMN_MAPPING, LOCATION_GROUPS,
    SPANISH_MONTHS
)



class ExcelService:
    """Service for handling Excel file operations."""
    
    def __init__(self, logger: logging.Logger):
        self.logger = logger
        # Callback for prompting user to select gender for ambiguous occupations
        # Signature: callback(person_name: str, cargo: str, male_option: str, female_option: str) -> str ('HOMBRE' or 'MUJER')
        self.gender_prompt_callback: Optional[callable] = None
    
    def load_excel_file(self, file_path: str) -> ExcelData:
        """
        Load Excel file and parse all worksheets.
        
        Args:
            file_path: Path to the Excel file
            
        Returns:
            ExcelData object with parsed worksheets or empty if failed
        """
        try:
            if not file_path:
                raise ValueError("File path is empty")
            
            if not Path(file_path).exists():
                raise FileNotFoundError(f"File does not exist: {file_path}")
            
            self.logger.info(f"Loading Excel file: {file_path}")
            
            # Read all sheets from Excel file
            with pd.ExcelFile(file_path) as excel_file:
                sheet_names = excel_file.sheet_names
                
                self.logger.info(f"Found {len(sheet_names)} worksheets: {sheet_names}")
                
                worksheets = []
                for sheet_name in sheet_names:
                    worksheet_result = self._parse_worksheet(excel_file, sheet_name)
                    worksheets.append(worksheet_result)
            
            excel_data = ExcelData(
                file_path=file_path,
                worksheets=worksheets
            )
            
            self.logger.info(f"Excel file loaded successfully. {excel_data.successful_worksheets}/{excel_data.total_worksheets} worksheets parsed, {excel_data.total_people_parsed} people total.")
            
            return excel_data
            
        except Exception as e:
            error_msg = f"Error loading Excel file: {str(e)}"
            self.logger.exception(error_msg)
            raise Exception(error_msg)
    
    def _parse_worksheet(self, excel_file: pd.ExcelFile, sheet_name: str) -> WorksheetParsingResult:
        """
        Parse a single worksheet and extract metadata and data.
        
        Args:
            excel_file: Pandas ExcelFile object
            sheet_name: Name of the sheet to parse
            
        Returns:
            WorksheetParsingResult with parsed data and metadata
        """
        metadata = WorksheetMetadata(sheet_name=sheet_name)
        result = WorksheetParsingResult(metadata=metadata)
        
        try:
            # Read the entire sheet first to extract metadata
            sheet_data = pd.read_excel(excel_file, sheet_name=sheet_name, header=None)
            
            # Extract metadata from specific cells
            try:
                if len(sheet_data) > METADATA_ROW_FECHA_SOLICITUD and len(sheet_data.columns) > METADATA_COL_FECHA_SOLICITUD:
                    metadata.fecha_solicitud = str(sheet_data.iloc[METADATA_ROW_FECHA_SOLICITUD, METADATA_COL_FECHA_SOLICITUD]) if pd.notna(sheet_data.iloc[METADATA_ROW_FECHA_SOLICITUD, METADATA_COL_FECHA_SOLICITUD]) else ""
                
                if len(sheet_data) > METADATA_ROW_TIENDA and len(sheet_data.columns) > METADATA_COL_TIENDA:
                    metadata.tienda = str(sheet_data.iloc[METADATA_ROW_TIENDA, METADATA_COL_TIENDA]) if pd.notna(sheet_data.iloc[METADATA_ROW_TIENDA, METADATA_COL_TIENDA]) else ""
                
                if len(sheet_data) > METADATA_ROW_ADMINISTRADOR and len(sheet_data.columns) > METADATA_COL_ADMINISTRADOR:
                    metadata.administrador = str(sheet_data.iloc[METADATA_ROW_ADMINISTRADOR, METADATA_COL_ADMINISTRADOR]) if pd.notna(sheet_data.iloc[METADATA_ROW_ADMINISTRADOR, METADATA_COL_ADMINISTRADOR]) else ""
                
                # Extract location group from LOCATION_ROW (row 6 in Excel = index 5)
                # This is the first non-empty cell in the location row that contains location group info
                if len(sheet_data) > LOCATION_ROW:
                    for col_idx in range(UNIFORM_DATA_START_COLUMN, min(UNIFORM_DATA_END_COLUMN + 1, len(sheet_data.columns))):
                        cell_value = sheet_data.iloc[LOCATION_ROW, col_idx]
                        if pd.notna(cell_value) and str(cell_value).strip():
                            metadata.location_group = str(cell_value).strip()
                            self.logger.info(f"Detected location group: '{metadata.location_group}' in sheet '{sheet_name}'")
                            break
            except Exception as meta_error:
                result.errors.append(f"Error extracting metadata: {str(meta_error)}")
                self.logger.warning(f"Metadata extraction error in sheet '{sheet_name}': {str(meta_error)}")
            
            # Extract main data (columns B through I) starting from configured row
            try:
                if len(sheet_data) > DATA_START_ROW:  # Ensure we have data rows
                    # Headers are in HEADER_ROW (row 8 in Excel = index 7)
                    # Data rows start at DATA_START_ROW (row 9 in Excel = index 8)
                    headers = sheet_data.iloc[HEADER_ROW, IGNORE_COLUMN_INDEX + 1:MAIN_DATA_END_COLUMN + 1]
                    
                    # Read main data: from DATA_START_ROW onwards, columns B through I
                    main_data_rows = sheet_data.iloc[DATA_START_ROW:, IGNORE_COLUMN_INDEX + 1:MAIN_DATA_END_COLUMN + 1]
                    main_data_rows.columns = headers
                    
                    if len(main_data_rows) > 0:
                        
                        # Remove completely empty rows from main data
                        main_data_rows = main_data_rows.dropna(how='all')
                        
                        # Check for missing DNI in rows with data
                        if 'DNI' in main_data_rows.columns or any('dni' in str(col).lower() for col in main_data_rows.columns):
                            # Find DNI column (case insensitive)
                            dni_col = None
                            for col in main_data_rows.columns:
                                if 'dni' in str(col).lower():
                                    dni_col = col
                                    break
                            
                            if dni_col is not None:
                                # Check for missing DNI in non-empty rows
                                non_empty_rows = main_data_rows.dropna(how='all')
                                missing_dni_count = non_empty_rows[dni_col].isna().sum()
                                
                                if missing_dni_count > 0:
                                    result.errors.append(f"{missing_dni_count} rows with data are missing DNI")
                        else:
                            result.warnings.append("No DNI column found - cannot validate DNI completeness")
                        
                        # Extract uniform data (columns J-BT, indices 9-71)
                        uniform_data = None
                        if len(sheet_data) > UNIFORM_DATA_START_ROW:
                            # Determine actual end column (some sheets may have fewer columns)
                            actual_end_col = min(UNIFORM_DATA_END_COLUMN + 1, len(sheet_data.columns))
                            uniform_data_rows = sheet_data.iloc[UNIFORM_DATA_START_ROW:, UNIFORM_DATA_START_COLUMN:actual_end_col]
                            
                            # Assign unique column names from mapping, matching actual column count
                            uniform_column_names = []
                            actual_col_count = len(uniform_data_rows.columns)
                            for i in range(actual_col_count):
                                col_idx = UNIFORM_DATA_START_COLUMN + i
                                if col_idx in UNIFORM_COLUMN_MAPPING:
                                    uniform_column_names.append(UNIFORM_COLUMN_MAPPING[col_idx])
                                else:
                                    uniform_column_names.append(f"PRENDA_{i}")
                            
                            uniform_data_rows.columns = uniform_column_names
                            
                            # CRITICAL: Join main and uniform data to perform aligned row cleaning
                            # This prevents uniform data from shifting if an employee has no uniform choices
                            combined_data = pd.concat([main_data_rows, uniform_data_rows], axis=1)
                            
                            # Drop rows where the person's name or logic identifier is missing
                            # We use APELLIDOS_Y_NOMBRES as it's the most reliable unique identifier for a row
                            if "APELLIDOS_Y_NOMBRES" in combined_data.columns:
                                combined_data = combined_data.dropna(subset=["APELLIDOS_Y_NOMBRES"])
                            else:
                                # Fallback if columns weren't named correctly for some reason
                                combined_data = combined_data.dropna(how='all')
                            
                            # Split back into main and uniform data
                            main_data_rows = combined_data.iloc[:, :len(main_data_rows.columns)]
                            uniform_data = combined_data.iloc[:, len(main_data_rows.columns):]
                            
    # Final reset of indices for both
                            main_data_rows = main_data_rows.reset_index(drop=True)
                            uniform_data = uniform_data.reset_index(drop=True)
                        else:
                            result.warnings.append("Insufficient data for uniform columns")
                        
                        result.data = main_data_rows
                        result.uniform_data = uniform_data
                        result.people_parsed = len(main_data_rows)  # After cleaning empty rows
                        result.total_lines = len(sheet_data)  # Total lines in sheet
                        
                        self.logger.info(f"Sheet '{sheet_name}': {result.people_parsed} people parsed from {result.total_lines} total lines")
                        if uniform_data is not None:
                            self.logger.info(f"Sheet '{sheet_name}': {len(uniform_data)} uniform data rows extracted")
                    else:
                        result.warnings.append("No data rows found after headers")
                else:
                    result.warnings.append(f"Sheet has insufficient rows (less than {DATA_START_ROW + 1})")
                    
            except Exception as data_error:
                result.errors.append(f"Error extracting data: {str(data_error)}")
                self.logger.error(f"Data extraction error in sheet '{sheet_name}': {str(data_error)}")
            
            # Validate metadata
            if not metadata.fecha_solicitud:
                result.warnings.append("Missing fecha_solicitud (C3)")
            if not metadata.tienda:
                result.warnings.append("Missing tienda (C4)")
            if not metadata.administrador:
                result.warnings.append("Missing administrador (C5)")
            
        except Exception as e:
            result.errors.append(f"Critical error parsing worksheet: {str(e)}")
            self.logger.exception(f"Critical error parsing worksheet '{sheet_name}': {str(e)}")
        
        return result
    
    def validate_excel_data(self, excel_data: ExcelData) -> ExcelValidationResult:
        """
        Validate Excel data for processing.
        
        Args:
            excel_data: ExcelData object to validate
            
        Returns:
            ExcelValidationResult with validation details
        """
        result = ExcelValidationResult(is_valid=False)
        
        try:
            if not excel_data.is_loaded:
                result.errors.append("No Excel worksheets could be loaded")
                result.message = "Excel file could not be loaded or has no valid worksheets"
                return result
            
            if excel_data.total_worksheets == 0:
                result.errors.append("Excel file contains no worksheets")
                result.message = "Excel file is empty or corrupted"
                return result
            
            # Validate each worksheet
            total_errors = 0
            total_warnings = 0
            
            for worksheet in excel_data.worksheets:
                if worksheet.errors:
                    result.errors.extend([f"Sheet '{worksheet.metadata.sheet_name}': {error}" for error in worksheet.errors])
                    total_errors += len(worksheet.errors)
                
                if worksheet.warnings:
                    result.warnings.extend([f"Sheet '{worksheet.metadata.sheet_name}': {warning}" for warning in worksheet.warnings])
                    total_warnings += len(worksheet.warnings)
            
            # Check if we have at least one successful worksheet
            if excel_data.successful_worksheets == 0:
                result.errors.append("No worksheets could be parsed successfully")
                result.message = f"All {excel_data.total_worksheets} worksheets failed to parse"
                return result
            
            # Check for occupation mapping issues
            occupation_mapping_issues = set()
            for worksheet in excel_data.worksheets:
                if worksheet.data is not None and 'cargo' in worksheet.data.columns:
                    for cargo in worksheet.data['cargo'].dropna().unique():
                        if cargo and str(cargo).strip():
                            normalized = self.unified_service.normalize_occupation(str(cargo))
                            if normalized != str(cargo).upper():
                                occupation_mapping_issues.add(f"'{cargo}' → '{normalized}'")
            
            # Add occupation mapping warnings
            if occupation_mapping_issues:
                result.warnings.append(f"Occupation mapping issues found: {', '.join(sorted(occupation_mapping_issues))}")
                result.warnings.append("Consider adding these occupations to the synonyms list in Configuration tab")
                total_warnings += len(occupation_mapping_issues) + 1
            
            # Validation passes if we have at least one successful worksheet
            result.is_valid = True
            
            # Create summary message
            success_msg = f"Excel file validated: {excel_data.successful_worksheets}/{excel_data.total_worksheets} worksheets parsed successfully"
            success_msg += f", {excel_data.total_people_parsed} people total"
            
            if total_errors > 0:
                success_msg += f", {total_errors} errors"
            if total_warnings > 0:
                success_msg += f", {total_warnings} warnings"
                
            result.message = success_msg
            
            return result
            
        except Exception as e:
            result.errors.append(f"Validation error: {str(e)}")
            result.message = "Unexpected error during validation"
            self.logger.exception(f"Excel validation error: {str(e)}")
            return result


class FileGenerationService:
    """Service for generating files from Excel data using templates."""
    
    def __init__(self, logger: logging.Logger, unified_service: UnifiedConfigService):
        self.logger = logger
        self.unified_service = unified_service
        # Callback for prompting user to select gender for ambiguous occupations
        self.gender_prompt_callback: Optional[callable] = None
    
    
    def generate_files(self, excel_data: ExcelData, config: AppConfig, options: Optional[GenerationOptions] = None) -> GenerationResult:
        """
        Generate files from Excel data using templates.
        
        Args:
            excel_data: Loaded Excel data
            config: Application configuration
            options: Generation options
            
        Returns:
            GenerationResult with success status and details
        """
        result = GenerationResult(success=False)
        
        try:
            # Validate inputs and setup
            validation_result = self._validate_generation_inputs(excel_data, config, options)
            if not validation_result.success:
                result.message = validation_result.message
                result.errors.extend(validation_result.errors)
                return result
            
            # Set default options if not provided
            if options is None:
                options = self._create_default_options(excel_data)
            
            # Group data by locale
            tienda_to_rows = self._group_data_by_locale(excel_data, options)
            if not tienda_to_rows:
                result.message = "No matching locales or people to generate"
                return result
            
            # Generate files
            files_generated = self._generate_documents(tienda_to_rows, config, options)
            
            # Create success result
            result.success = True
            result.files_generated = files_generated
            total_people = sum(len(people) for people in tienda_to_rows.values())
            templates = self._get_enabled_templates(options)
            result.message = f"Generated {', '.join(templates)} documents for {total_people} people across {len(tienda_to_rows)} locales"
            
            self.logger.info(result.message)
            return result
            
        except Exception as e:
            error_msg = f"Error during file generation: {str(e)}"
            self.logger.exception(error_msg)
            result.message = error_msg
            result.errors.append(error_msg)
            return result

    def _validate_generation_inputs(self, excel_data: ExcelData, config: AppConfig, options: Optional[GenerationOptions]) -> GenerationResult:
        """Validate inputs for file generation."""
        result = GenerationResult(success=False)
        
        if not excel_data.is_loaded:
            result.message = "No Excel data loaded"
            return result
        
        # Validate template files
        template_errors = TemplateValidator.validate_autorizacion_template(config)
        if template_errors:
            result.message = "Template validation failed"
            result.errors.extend(template_errors)
            return result
        
        # Create and validate destination directory
        dest_path = Path(config.destination_path)
        dest_path.mkdir(parents=True, exist_ok=True)
        if not dest_path.exists() or not dest_path.is_dir():
            result.message = f"Destination path is invalid: {dest_path}"
            return result
        
        # Check if any templates are enabled
        if options and not (options.autorizacion_enabled or options.cargo_enabled):
            result.message = "No templates selected for generation"
            return result
        
        # Check docxtpl availability
        try:
            from docxtpl import DocxTemplate  # noqa: F401
        except Exception as e:
            result.message = "docxtpl not available; please install python-docx-template"
            result.errors.append(str(e))
            return result
        
        result.success = True
        return result
    
    def _create_default_options(self, excel_data: ExcelData) -> GenerationOptions:
        """Create default generation options."""
        locales = [w.metadata.tienda for w in excel_data.worksheets if w.metadata.tienda]
        return GenerationOptions(
            selected_locales=locales,
            combine_per_local=False,
            cargo_enabled=True,
            autorizacion_enabled=True
        )
    
    def _get_enabled_templates(self, options: GenerationOptions) -> List[str]:
        """Get list of enabled template names."""
        templates = []
        if options.autorizacion_enabled:
            templates.append("AUTORIZACION")
        if options.cargo_enabled:
            templates.append("CARGO")
        return templates
    
    def _group_data_by_locale(self, excel_data: ExcelData, options: GenerationOptions) -> Dict[str, List[Dict[str, Any]]]:
        """Group worksheet data by locale (tienda) with context building."""
        tienda_to_rows: Dict[str, List[Dict[str, Any]]] = {}
        
        for ws in excel_data.worksheets:
            tienda = str(ws.metadata.tienda or "").strip()
            if not tienda or tienda not in options.selected_locales:
                continue
            
            if ws.data is None or ws.data.empty:
                continue
            
            # Skip sheets with critical errors (missing fecha_solicitud, tienda, or DNI errors)
            # Skip only if tienda is missing (critical for folder structure)
            if not tienda or tienda.strip() == "":
                self.logger.warning(f"Sheet '{ws.metadata.sheet_name}' skipped: Missing tienda")
                continue
            
            # Process each person in the worksheet
            for idx, row in ws.data.iterrows():
                person_contexts = self._build_person_contexts(row, ws, options)
                if person_contexts:
                    tienda_to_rows.setdefault(tienda, []).append(person_contexts)
        
        return tienda_to_rows
    
    
    
    def _build_person_contexts(self, row: pd.Series, ws, options: GenerationOptions) -> Dict[str, Any]:
        """Build context dictionaries for a person based on enabled templates."""
        person_contexts = {}
        
        # Get uniform row for both contexts
        uniform_row = self._get_uniform_row_for_person(row, ws)
        
        if options.autorizacion_enabled:
            autorizacion_ctx = self._build_autorizacion_context(row, ws.metadata, uniform_row)
            if autorizacion_ctx:
                person_contexts["AUTORIZACION"] = autorizacion_ctx
                self.logger.debug(f"Built AUTORIZACION context for {autorizacion_ctx.get('nombre', 'unknown')}")
            else:
                self.logger.warning("Failed to build AUTORIZACION context for row")
        
        if options.cargo_enabled:
            cargo_ctx = self._build_cargo_context(row, ws.metadata, uniform_row)
            if cargo_ctx:
                person_contexts["CARGO"] = cargo_ctx
                self.logger.debug(f"Built CARGO context for {cargo_ctx.get('nombre', 'unknown')} with {len(cargo_ctx.get('prendas', []))} prendas")
            else:
                self.logger.warning("Failed to build CARGO context for row")
        
        return person_contexts
    
    def _generate_documents(self, tienda_to_rows: Dict[str, List[Dict[str, Any]]], config: AppConfig, options: GenerationOptions) -> int:
        """Generate all documents for all locales and people."""
        dest_path = Path(config.destination_path)
        files_generated = 0
        
        for tienda, people in tienda_to_rows.items():
            tienda_folder = dest_path / self._sanitize_name(tienda)
            tienda_folder.mkdir(parents=True, exist_ok=True)
            
            template_docs: Dict[str, List[Path]] = {}
            
            # Generate individual documents for each person
            for person_contexts in people:
                person_name = self._extract_person_name(person_contexts)
                if not person_name:
                    self.logger.warning("Skipping person with no name")
                    continue
                
                person_folder = tienda_folder / self._sanitize_name(person_name)
                person_folder.mkdir(parents=True, exist_ok=True)
                
                # Generate documents for each template type
                for template_type, context in person_contexts.items():
                    docx_path = self._generate_single_document(
                        template_type, context, person_folder, config
                    )
                    if docx_path:
                        template_docs.setdefault(template_type, []).append(docx_path)
                        files_generated += 1
            
            # Create combined documents if requested
            if options.combine_per_local:
                files_generated += self._create_combined_documents(template_docs, tienda_folder, tienda)
        
        return files_generated
    
    def _extract_person_name(self, person_contexts: Dict[str, Any]) -> str:
        """Extract person name from any available context."""
        for context in person_contexts.values():
            if context.get("nombre"):
                return context["nombre"]
        return ""
    
    def _generate_single_document(self, template_type: str, context: Dict[str, Any], person_folder: Path, config: AppConfig) -> Optional[Path]:
        """Generate a single document for a person."""
        try:
            if template_type == "AUTORIZACION":
                docx_path = person_folder / f"AUTORIZACION_{self._file_stub(context)}.docx"
                self.logger.info(f"Generating AUTORIZACION document: {docx_path}")
                self._render_autorizacion_doc(config.autorizacion_template_path, context, docx_path)
                return docx_path
            
            elif template_type == "CARGO":
                docx_path = person_folder / f"CARGO_{self._file_stub(context)}.docx"
                self.logger.info(f"Generating CARGO document: {docx_path}")
                self._render_cargo_doc(config.cargo_template_path, context, docx_path)
                return docx_path
            
        except Exception as e:
            self.logger.error(f"Failed to generate {template_type} document: {e}")
        
        return None
    
    def _create_combined_documents(self, template_docs: Dict[str, List[Path]], tienda_folder: Path, tienda: str) -> int:
        """Create combined documents for each template type."""
        files_generated = 0
        
        for template_type, docs in template_docs.items():
            if docs:
                combined_path = tienda_folder / f"{template_type}_COMBINED_{self._sanitize_name(tienda)}.docx"
                self._create_combined_docx(docs, combined_path)
                files_generated += 1
        
        return files_generated

    def _extract_common_person_data(self, row: pd.Series) -> Dict[str, str]:
        """Extract common person data from a row."""
        return {
            "cargo": self._find_in_row(row, ["cargo"]) or "",
            "nombre": self._find_in_row(row, ["nombre"]) or self._extract_name(row),
            "identificacion": self._find_in_row(row, ["dni"]) or ""
        }
    
    def _get_uniform_row_for_person(self, row: pd.Series, ws) -> Optional[pd.Series]:
        """Get uniform row for a person, handling edge cases safely."""
        if not self._has_valid_uniform_data(ws):
            return None
        
        row_index = row.name if hasattr(row, 'name') else 0
        if row_index < len(ws.uniform_data):
            return ws.uniform_data.iloc[row_index]
        return None
    
    def _has_valid_uniform_data(self, ws) -> bool:
        """Check if worksheet has valid uniform data."""
        return (ws.uniform_data is not None and 
                hasattr(ws, 'data') and 
                len(ws.data) > 0)
    
    def _calculate_juegos(self, row: pd.Series, metadata: WorksheetMetadata, uniform_row: Optional[pd.Series] = None) -> int:
        """Calculate the juegos value based on primary prenda or mode of other prendas."""
        try:
            # Get cargo and normalize it
            cargo = self._find_in_row(row, ["cargo"]) or ""
            person_name = self._extract_name(row)
            self._current_person_name = person_name
            
            if not cargo:
                self.logger.warning(f"No cargo found for person {person_name}, cannot calculate juegos")
                return 0
            
            # Handle gendered occupations - detect gender from prenda columns
            if self._is_gendered_occupation(cargo):
                data_row = uniform_row if uniform_row is not None else row
                detected_gender = self._detect_gender_from_row(data_row)
                cargo = self._resolve_gendered_occupation(cargo, detected_gender)
            
            # Normalize cargo using synonyms
            normalized_cargo = self.unified_service.normalize_occupation(cargo)
            person_name = self._extract_name(row)
            self.logger.debug(f"Calculating juegos for {person_name}: cargo='{cargo}' -> normalized='{normalized_cargo}'")
            
            # Get primary prenda for this occupation
            primary_prenda_cfg = self.unified_service.get_primary_prenda(normalized_cargo)
            if not primary_prenda_cfg:
                self.logger.warning(f"No primary prenda configured for occupation: {normalized_cargo} (person: {person_name})")
                return 0
            
            primary_prenda_type = primary_prenda_cfg.prenda_type.upper().strip()
            self.logger.debug(f"Primary prenda for {normalized_cargo}: {primary_prenda_type}")
            
            # Build prendas list (same as in _get_monto_for_person)
            talla_superior = self._extract_talla_superior(row)
            prendas = self._build_prendas_list(uniform_row if uniform_row is not None else row, talla_superior)
            
            if not prendas:
                self.logger.warning(f"No prendas found for person {person_name}, cannot calculate juegos")
                return 0
            
            self.logger.debug(f"Found {len(prendas)} prendas for {person_name}: {[p.get('prenda_type', '') + ' (qty=' + str(p.get('qty', 0)) + ')' for p in prendas]}")
            
            # Look for primary prenda in the prendas list
            primary_prenda_found = None
            
            for prenda in prendas:
                prenda_type = prenda.get("prenda_type", "").upper().strip()
                if prenda_type == primary_prenda_type:
                    primary_prenda_found = prenda
                    break
            
            # If primary prenda found and qty > 0, use that
            if primary_prenda_found:
                qty = primary_prenda_found.get("qty", 0)
                if qty > 0:
                    juegos = int(qty)
                    self.logger.info(f"Juegos calculated from primary prenda {primary_prenda_type} for {person_name}: {juegos}")
                    return juegos
                else:
                    self.logger.debug(f"Primary prenda {primary_prenda_type} found but qty=0 for {person_name}, calculating mode of other prendas")
            else:
                self.logger.debug(f"Primary prenda {primary_prenda_type} not found in prendas list for {person_name}, calculating mode of other prendas")
            
            # If primary prenda not found or qty = 0, calculate mode of other prendas
            other_quantities = []
            for prenda in prendas:
                prenda_type = prenda.get("prenda_type", "").upper().strip()
                qty = prenda.get("qty", 0)
                # Skip primary prenda and only include prendas with qty > 0
                if prenda_type != primary_prenda_type and qty > 0:
                    other_quantities.append(int(qty))
            
            if not other_quantities:
                self.logger.warning(f"No other prendas with qty > 0 found for {person_name}, returning 0 for juegos")
                return 0
            
            # Calculate mode (most common value)
            try:
                from statistics import mode, StatisticsError
                juegos = int(mode(other_quantities))
            except (ImportError, StatisticsError):
                # Fallback: manual mode calculation
                from collections import Counter
                quantity_counts = Counter(other_quantities)
                # Get the most common value(s)
                max_count = max(quantity_counts.values())
                modes = [qty for qty, count in quantity_counts.items() if count == max_count]
                # If tie, return the highest value
                juegos = int(max(modes)) if modes else 0
            
            self.logger.info(f"Juegos calculated from mode of other prendas for {person_name}: {juegos} (from quantities: {other_quantities})")
            return juegos
            
        except Exception as e:
            person_name = self._extract_name(row) if hasattr(self, '_extract_name') else "unknown"
            self.logger.error(f"Failed to calculate juegos for person {person_name}: {e}", exc_info=True)
            return 0
    
    def _build_autorizacion_context(self, row: pd.Series, metadata: WorksheetMetadata, uniform_row: Optional[pd.Series] = None) -> Optional[Dict[str, Any]]:
        """Build docxtpl context for AUTORIZACION: dia, mes, anho, local, cargo, nombre, identificacion, monto, juegos."""
        try:
            # Parse fecha_solicitud using the same flexible parsing as CARGO
            dia, _, anho, fecha_formatted = self._get_system_date()
            
            # Extract numeric month for AUTORIZACION (it expects MM format)
            from datetime import datetime
            now = datetime.now()
            mes = f"{now.month:02d}"
            
            # Extract common person data
            person_data = self._extract_common_person_data(row)
            
            # Format fecha as dd / MM / yyyy for AUTORIZACION template
            fecha_template = ""
            if dia and mes and anho:
                fecha_template = f"{dia} / {mes} / {anho}"
            
            # Get monto (calculated from pricing service)
            monto_value = self._get_monto_for_person(row, metadata, uniform_row)
            monto_formatted = f"S/ {monto_value:.2f}"
            
            # Calculate juegos (based on primary prenda or mode of other prendas)
            try:
                juegos_value = self._calculate_juegos(row, metadata, uniform_row)
                # Ensure juegos is always an integer (never None)
                juegos_value = int(juegos_value) if juegos_value is not None else 0
            except Exception as e:
                self.logger.error(f"Error calculating juegos: {e}", exc_info=True)
                juegos_value = 0  # Default to 0 if calculation fails
            
            # Debug logging for missing data
            if not person_data["identificacion"]:
                self.logger.warning(f"No DNI found for person: {person_data['nombre']}. Available columns: {list(row.index)}")
            if not person_data["nombre"]:
                self.logger.warning(f"No nombre found. Available columns: {list(row.index)}")
            if not person_data["cargo"]:
                self.logger.warning(f"No cargo found. Available columns: {list(row.index)}")
            if not fecha_template:
                self.logger.warning(f"No date found in metadata: {metadata.fecha_solicitud}")
            
            context = {
                "dia": dia,
                "mes": mes,
                "anho": anho,
                "fecha": fecha_template,  # Add formatted date for template
                "local": metadata.tienda or "",
                "cargo": str(person_data["cargo"]),
                "nombre": str(person_data["nombre"]),
                "identificacion": str(person_data["identificacion"]),
                "monto": monto_formatted,  # Formatted with Sol currency
                "juegos": juegos_value,  # Number of juegos (calculated from primary prenda or mode) - always an integer
            }
            
            # Log the context to verify juegos is included
            self.logger.info(f"AUTORIZACION context for {person_data['nombre']}: juegos={juegos_value} (type: {type(juegos_value).__name__}), cargo={person_data['cargo']}, local={metadata.tienda}")
            self.logger.debug(f"Full AUTORIZACION context keys: {list(context.keys())}")
            self.logger.debug(f"Full AUTORIZACION context: {context}")
            
            return context
        except Exception as e:
            self.logger.error(f"Error building context: {e}")
            return None

    def _build_cargo_context(self, row: pd.Series, metadata: WorksheetMetadata, uniform_row: Optional[pd.Series] = None) -> Optional[Dict[str, Any]]:
        """Build docxtpl context for CARGO documents with Spanish months and prenda handling."""
        try:
            # Parse fecha_solicitud with flexible date handling
            dia, mes_string, anho, fecha_string = self._get_system_date()
            
            # Extract common person data
            person_data = self._extract_common_person_data(row)
            
            # Extract talla prenda superior (second to last row data)
            talla_superior = self._extract_talla_superior(row)
            
            # Build prendas list from uniform data
            # Use uniform_row if available, otherwise fall back to main row
            data_row = uniform_row if uniform_row is not None else row
            prendas = self._build_prendas_list(data_row, talla_superior)
            
            # Get monto (calculated from pricing service)
            monto_value = self._get_monto_for_person(row, metadata, uniform_row)
            monto_formatted = f"S/ {monto_value:.2f}"
            
            # Calculate juegos (based on primary prenda or mode of other prendas)
            try:
                juegos_value = self._calculate_juegos(row, metadata, uniform_row)
                # Ensure juegos is always an integer (never None)
                juegos_value = int(juegos_value) if juegos_value is not None else 0
            except Exception as e:
                self.logger.error(f"Error calculating juegos: {e}", exc_info=True)
                juegos_value = 0  # Default to 0 if calculation fails
            
            context = {
                "dia": dia,
                "mes_string": mes_string,
                "anho": anho,
                "fecha": fecha_string,
                "nombre": str(person_data["nombre"]),
                "prendas": prendas,
                "monto": monto_formatted,  # Add pricing information
                "juegos": juegos_value,  # Number of juegos (calculated from primary prenda or mode) - always an integer
            }
            
            return context
        except Exception as e:
            self.logger.error(f"Error building CARGO context: {e}")
            return None

    def _render_autorizacion_doc(self, template_path: str, context: Dict[str, Any], output_docx: Path) -> None:
        """Render AUTORIZACION document."""
        self._render_document(template_path, context, output_docx)
    
    def _render_cargo_doc(self, template_path: str, context: Dict[str, Any], output_docx: Path) -> None:
        """Render CARGO document."""
        self._render_document(template_path, context, output_docx)
    
    def _render_document(self, template_path: str, context: Dict[str, Any], output_docx: Path) -> None:
        """Generic document rendering method."""
        from docxtpl import DocxTemplate
        # Log context before rendering to verify all variables are present
        self.logger.debug(f"Rendering document {output_docx.name} with context keys: {list(context.keys())}")
        if "juegos" in context:
            self.logger.info(f"Context includes 'juegos' = {context['juegos']} (type: {type(context['juegos'])})")
        else:
            self.logger.warning(f"Context does NOT include 'juegos'! Available keys: {list(context.keys())})")
        tpl = DocxTemplate(template_path)
        tpl.render(context)
        tpl.save(str(output_docx))

    def _create_combined_docx(self, individual_docs: List[Path], output_path: Path) -> None:
        """Combine multiple DOCX files into one document with proper formatting using docxcompose."""
        try:
            from docxcompose.composer import Composer
            from docx import Document
            
            if not individual_docs:
                self.logger.warning("No individual documents to combine")
                return
            
            # Filter out non-existent documents
            valid_docs = [doc for doc in individual_docs if doc.exists()]
            if not valid_docs:
                self.logger.error("No valid documents to combine")
                return
            
            # Load the first document as the master
            master_doc = Document(str(valid_docs[0]))
            composer = Composer(master_doc)
            
            # Append each subsequent document with a page break before it
            for doc_path in valid_docs[1:]:
                try:
                    # Add a page break at the end of the current master document
                    # This ensures the next document starts on a new page
                    master_doc.add_page_break()
                    
                    doc_to_append = Document(str(doc_path))
                    composer.append(doc_to_append)
                    self.logger.debug(f"Successfully appended: {doc_path}")
                except Exception as e:
                    self.logger.warning(f"Failed to append document {doc_path}: {e}")
                    continue
            
            # Save the merged document
            composer.save(str(output_path))
            self.logger.info(f"Created combined document with preserved formatting: {output_path}")
            
        except ImportError as e:
            self.logger.error(f"docxcompose not available: {e}")
            # Fallback to manual method
            self._create_fallback_combined_docx(individual_docs, output_path)
            
        except Exception as e:
            self.logger.error(f"Failed to create combined document with docxcompose: {e}")
            # Fallback to manual method
            self._create_fallback_combined_docx(individual_docs, output_path)
    
    def _create_fallback_combined_docx(self, individual_docs: List[Path], output_path: Path) -> None:
        """Fallback method for combining documents when docxcompose fails."""
        try:
            from docx import Document
            
            # Create a simple document listing all individual files
            fallback_doc = Document()
            fallback_doc.add_heading(f"Combined Document - {output_path.stem}", 0)
            fallback_doc.add_paragraph(f"This document combines {len(individual_docs)} individual authorization documents.")
            fallback_doc.add_paragraph("Individual documents:")
            
            for i, doc_path in enumerate(individual_docs, 1):
                if doc_path.exists():
                    fallback_doc.add_paragraph(f"{i}. {doc_path.name}")
                else:
                    fallback_doc.add_paragraph(f"{i}. {doc_path.name} (missing)")
            
            fallback_doc.add_paragraph("")
            fallback_doc.add_paragraph("Note: Please check individual documents for complete content.")
            fallback_doc.save(str(output_path))
            self.logger.info(f"Created fallback combined document: {output_path}")
            
        except Exception as e:
            self.logger.error(f"Failed to create fallback combined document: {e}")
    
    def _get_monto_for_person(self, row: pd.Series, metadata: WorksheetMetadata, uniform_row: Optional[pd.Series] = None) -> float:
        """Get the monto (amount) for a person based on their uniform requirements."""
        try:
            # Get cargo and normalize it
            cargo = self._find_in_row(row, ["cargo"]) or ""
            if not cargo:
                self.logger.warning("No cargo found for person, using default pricing")
                cargo = "MOZO"  # Default cargo
            
            # Handle gendered occupations - detect gender from prenda columns
            if self._is_gendered_occupation(cargo):
                data_row = uniform_row if uniform_row is not None else row
                detected_gender = self._detect_gender_from_row(data_row)
                # Store person name for callback context
                self._current_person_name = self._extract_name(row)
                cargo = self._resolve_gendered_occupation(cargo, detected_gender)
            
            # Normalize cargo using synonyms
            normalized_cargo = self.unified_service.normalize_occupation(cargo)
            
            # Check if occupation was properly mapped - show detailed info for gendered occupations
            if normalized_cargo != cargo.upper():
                # Get a sample price for the mapped occupation to show user
                occupation_obj = self.unified_service.get_occupation(normalized_cargo)
                sample_price = ""
                if occupation_obj and occupation_obj.prendas:
                    first_prenda = occupation_obj.prendas[0]
                    sample_price = f" (sample: {first_prenda.display_name} S/{first_prenda.price_sml_other})"
                
                # Show clear alert for gendered mappings (ANFITRIONAJE, CAJA, STAFF, etc.)
                if "(HOMBRE)" in normalized_cargo or "(MUJER)" in normalized_cargo:
                    gender = "👨 MALE" if "(HOMBRE)" in normalized_cargo else "👩 FEMALE"
                    self.logger.info(f"🔔 GENDER SELECTION: '{cargo}' → '{normalized_cargo}' [{gender}]{sample_price}")
                else:
                    self.logger.warning(f"⚠️  OCCUPATION MAPPING: '{cargo}' → '{normalized_cargo}'{sample_price} - consider adding '{cargo}' to synonyms")
            else:
                # Direct match - still log for visibility
                occupation_obj = self.unified_service.get_occupation(normalized_cargo)
                if occupation_obj and occupation_obj.prendas:
                    first_prenda = occupation_obj.prendas[0]
                    self.logger.info(f"✓ Occupation: '{normalized_cargo}' (sample: {first_prenda.display_name} S/{first_prenda.price_sml_other})")
            
            # Get local/tienda
            local = metadata.tienda or "OTHER"
            
            # Build prendas list for pricing calculation
            talla_superior = self._extract_talla_superior(row)
            prendas = self._build_prendas_list(uniform_row if uniform_row is not None else row, talla_superior)
            
            # Debug: log detailed information for PACKER and MOTORIZADO
            person_name = self._extract_name(row)
            if normalized_cargo in ['PACKER', 'MOTORIZADO'] or 'PACKER' in cargo.upper() or 'MOTORIZADO' in cargo.upper():
                self.logger.info(f"DEBUG PACKER/MOTORIZADO - {person_name}:")
                self.logger.info(f"  Original cargo: '{cargo}'")
                self.logger.info(f"  Normalized cargo: '{normalized_cargo}'")
                self.logger.info(f"  Local: '{local}'")
                self.logger.info(f"  Talla superior: '{talla_superior}'")
                self.logger.info(f"  Uniform row available: {uniform_row is not None}")
                if uniform_row is not None:
                    self.logger.info(f"  Uniform row columns: {list(uniform_row.index)}")
                    self.logger.info(f"  Uniform row values: {uniform_row.to_dict()}")
                self.logger.info(f"  Found {len(prendas)} prendas:")
                for i, prenda in enumerate(prendas):
                    self.logger.info(f"    {i+1}. {prenda}")
            
            # Calculate total price
            total_price = self.unified_service.calculate_total_price(
                prendas, normalized_cargo, local, local_group=metadata.location_group
            )
            
            self.logger.info(f"Calculated monto for {person_name}: "
                           f"Cargo={normalized_cargo}, Local={local}, Prendas={len(prendas)}, Total={total_price}")
            
            # Debug: log individual prenda prices
            for prenda in prendas:
                self.logger.debug(f"  Prenda: {prenda['string']}, Qty: {prenda['qty']}")
            
            return total_price
            
        except Exception as e:
            self.logger.error(f"Failed to calculate monto for person: {e}")
            return 0.0
    
    def _get_system_date(self) -> tuple:
        """Get system date in the required format."""
        from datetime import datetime
        now = datetime.now()
        dia = f"{now.day:02d}"
        mes_string = SPANISH_MONTHS.get(now.month, "")
        anho = f"{now.year}"
        fecha_string = f"{dia} de {mes_string} de {anho}"
        return dia, mes_string, anho, fecha_string
    
    def _extract_talla_superior(self, row: pd.Series) -> str:
        """Extract talla prenda superior from row data."""
        # Look for talla-related columns (second to last priority)
        talla_keys = ["talla prenda superior", "talla superior", "talla", "size", "talla_superior"]
        talla = self._find_in_row(row, talla_keys) or ""
        return str(talla).strip().upper()
    
    def _extract_talla_inferior(self, row: pd.Series) -> str:
        """Extract talla prenda inferior from row data."""
        # Look for talla inferior columns
        talla_keys = ["talla prenda inferior", "talla inferior", "talla_inferior"]
        talla = self._find_in_row(row, talla_keys) or ""
        return str(talla).strip().upper()
    
    def _get_talla_for_garment(self, row: pd.Series, garment_type: str) -> str:
        """Get the appropriate talla based on garment type (UPPER or LOWER)."""
        if garment_type == "LOWER":
            # Try to get talla inferior first
            talla_inferior = self._extract_talla_inferior(row)
            if talla_inferior:
                return talla_inferior
        
        # Default to talla superior (for UPPER garments or as fallback)
        return self._extract_talla_superior(row)
    
    def _determine_garment_type(self, prenda_type: str) -> str:
        """Determine if a prenda is UPPER or LOWER body garment."""
        prenda_upper = prenda_type.upper()
        
        # Lower body garments
        lower_garments = {"PANTALON", "PANTALÓN", "PANTS", "TROUSERS"}
        if any(lower in prenda_upper for lower in lower_garments):
            return "LOWER"
        
        # Everything else defaults to UPPER (shirts, jackets, aprons, etc.)
        return "UPPER"
    
    def _build_prendas_list(self, row: pd.Series, talla_superior: str) -> List[Dict[str, Any]]:
        """Build list of prendas from uniform data with quantities."""
        prendas = []
        
        # Enhanced approach: Check all uniform columns dynamically
        uniform_columns = self._get_uniform_columns_from_row(row)
        
        for column_name in uniform_columns:
            qty_value = row[column_name]
            
            # Handle Series values (from duplicate column names)
            if isinstance(qty_value, pd.Series):
                qty_value = qty_value.iloc[0] if len(qty_value) > 0 else None
            
            # Check if we have a valid quantity value (not NaN, not empty, not zero)
            if (qty_value is not None and 
                pd.notna(qty_value) and 
                str(qty_value).strip() not in ['', 'nan', 'NaN', '0']):
                try:
                    qty = int(float(str(qty_value).strip()))
                    if qty > 0:
                        # Normalize column name to prenda type
                        prenda_type = self._normalize_prenda_type(column_name)
                        display_name = self._get_display_name(prenda_type)
                        
                        # Determine garment type and get appropriate talla
                        garment_type = self._determine_garment_type(prenda_type)
                        talla = self._get_talla_for_garment(row, garment_type)
                        
                        # Use the determined talla (not just talla_superior)
                        if not talla:
                            talla = talla_superior  # Fallback to talla_superior if not found
                        
                        # Create formatted prenda string for display
                        prenda_string = self._format_prenda_string(display_name, talla)
                        
                        prenda_dict = {
                            "string": prenda_string,
                            "qty": qty,
                            "prenda_type": prenda_type,
                            "garment_type": garment_type,
                            "talla": talla
                        }
                        prendas.append(prenda_dict)
                        
                except (ValueError, TypeError):
                    # Only log warning for truly invalid values, not NaN/empty
                    if str(qty_value).strip() not in ['', 'nan', 'NaN']:
                        self.logger.warning(f"Invalid quantity value for {column_name}: {qty_value}")
                    continue
        prendas = self._apply_business_rules(prendas, row)
        
        return prendas
    
    def _apply_business_rules(self, prendas: List[Dict[str, Any]], row: pd.Series) -> List[Dict[str, Any]]:
        """Apply business rules to prendas list.
        
        Rules:
        1. All male admin (STAFF ADMINISTRATIVO HOMBRE, ADMINISTRADOR) get 2 corbatas
        2. SACO quantity is always max 1 for any occupation
        """
        # Get occupation from row to determine if male admin
        cargo = self._find_in_row(row, ["cargo"]) or ""
        cargo_upper = cargo.upper().strip()
        
        # Determine if this is a male admin role
        is_male_admin = any(x in cargo_upper for x in [
            "STAFF ADMINISTRATIVO (HOMBRE)",
            "STAFF ADMINISTRATIVO (H)",
            "ADMINISTRADOR (H)",
            "ADMINISTRADOR (HOMBRE)",
        ]) or (
            "ADMIN" in cargo_upper and 
            "(M)" not in cargo_upper and 
            "MUJER" not in cargo_upper and
            "(F)" not in cargo_upper
        )
        
        # Rule 1: Male admin gets 2 corbatas
        if is_male_admin:
            has_corbata = any(p.get('prenda_type', '').upper() == 'CORBATA' for p in prendas)
            if not has_corbata:
                # Add corbata
                prendas.append({
                    "string": "Corbata",
                    "qty": 2,
                    "prenda_type": "CORBATA",
                    "garment_type": "ACCESSORY",
                    "talla": ""
                })
                self.logger.info(f"Business rule: Added 2 CORBATA for male admin '{cargo}'")
            else:
                # Force quantity to 2
                for p in prendas:
                    if p.get('prenda_type', '').upper() == 'CORBATA':
                        if p.get('qty', 0) != 2:
                            self.logger.info(f"Business rule: Adjusted CORBATA qty from {p.get('qty')} to 2 for '{cargo}'")
                            p['qty'] = 2
        
        # Rule 2: SACO always max 1
        for p in prendas:
            if p.get('prenda_type', '').upper() == 'SACO' and p.get('qty', 0) > 1:
                self.logger.info(f"Business rule: Capped SACO qty from {p.get('qty')} to 1 for '{cargo}'")
                p['qty'] = 1
        
        return prendas
    
    def _detect_gender_from_row(self, row: pd.Series) -> Optional[str]:
        """Detect gender from prenda columns.
        
        Returns:
            'HOMBRE' - if male-specific prendas found (CAMISA, SACO_H)
            'MUJER' - if female-specific prendas found (BLUSA, SACO_M)
            None - if gender cannot be determined
        """
        # Check for male indicators
        male_indicators = ['CAMISA', 'SACO_H', 'SACO H']
        female_indicators = ['BLUSA', 'SACO_M', 'SACO M']
        
        has_male = False
        has_female = False
        
        for col in row.index:
            col_upper = str(col).upper()
            value = row[col]
            
            # Check if column has a positive quantity
            has_value = False
            if pd.notna(value):
                try:
                    qty = int(float(str(value).strip()))
                    has_value = qty > 0
                except (ValueError, TypeError):
                    pass
            
            if has_value:
                if any(ind in col_upper for ind in male_indicators):
                    has_male = True
                elif any(ind in col_upper for ind in female_indicators):
                    has_female = True
        
        if has_male and not has_female:
            return 'HOMBRE'
        elif has_female and not has_male:
            return 'MUJER'
        return None
    
    def _is_gendered_occupation(self, cargo: str) -> bool:
        """Check if occupation name is ambiguous/gendered (contains (A) or similar)."""
        cargo_upper = cargo.upper().strip()
        return any(pattern in cargo_upper for pattern in [
            '(A)', '(O/A)', '(A/O)', '(O)', '(HOMBRE/MUJER)', '(MUJER/HOMBRE)'
        ])
    
    def _resolve_gendered_occupation(self, cargo: str, detected_gender: Optional[str]) -> str:
        """Resolve a gendered occupation to its specific variant.
        
        Args:
            cargo: Original cargo from Excel (e.g., "ADMINISTRADOR(A)")
            detected_gender: 'HOMBRE' or 'MUJER' or None
        
        Returns:
            The resolved occupation name
        """
        cargo_upper = cargo.upper().strip()
        
        # Common gendered occupation patterns
        gendered_patterns = {
            'ADMINISTRADOR(A)': ('STAFF ADMINISTRATIVO (HOMBRE)', 'STAFF ADMINISTRATIVO (MUJER)'),
            'ADMINISTRADOR (A)': ('STAFF ADMINISTRATIVO (HOMBRE)', 'STAFF ADMINISTRATIVO (MUJER)'),
            'CAJERO(A)': ('CAJA (HOMBRE)', 'CAJA (MUJER)'),
            'CAJERO (A)': ('CAJA (HOMBRE)', 'CAJA (MUJER)'),
            'MOZO(A)': ('MOZO', 'MOZA'),  # If these exist
        }
        
        for pattern, (male_occ, female_occ) in gendered_patterns.items():
            if pattern in cargo_upper:
                if detected_gender == 'MUJER':
                    self.logger.info(f"Resolved gendered occupation: {cargo} -> {female_occ}")
                    return female_occ
                elif detected_gender == 'HOMBRE':
                    self.logger.info(f"Resolved gendered occupation: {cargo} -> {male_occ}")
                    return male_occ
                else:
                    # Gender undetermined - prompt user if callback is set
                    if self.gender_prompt_callback:
                        person_name = self._current_person_name if hasattr(self, '_current_person_name') else 'Unknown'
                        
                        # Check cache primarily by person_name + cargo to avoid reprompting for same person
                        cache_key = f"{person_name}|{cargo}"
                        if not hasattr(self, '_gender_selection_cache'):
                            self._gender_selection_cache = {}
                        
                        if cache_key in self._gender_selection_cache:
                            selected = self._gender_selection_cache[cache_key]
                            self.logger.info(f"Using cached gender selection for {person_name}: {selected}")
                            return female_occ if selected == 'MUJER' else male_occ

                        selected = self.gender_prompt_callback(person_name, cargo, male_occ, female_occ)
                        
                        # Cache the selection
                        self._gender_selection_cache[cache_key] = selected
                        
                        if selected == 'MUJER':
                            self.logger.info(f"User selected MUJER for {cargo} -> {female_occ}")
                            return female_occ
                        else:
                            self.logger.info(f"User selected HOMBRE for {cargo} -> {male_occ}")
                            return male_occ
                    else:
                        # No callback - log warning and default to male (fallback)
                        self.logger.warning(f"Gender undetermined for '{cargo}' - no prompt callback, defaulting to male")
                        return male_occ
        
        # If no pattern matched, return as-is
        return cargo
    
    def _get_uniform_columns_from_row(self, row: pd.Series) -> List[str]:
        """Get all uniform-related columns from a row dynamically."""
        uniform_columns = []
        
        # Get all columns that could be uniform data
        for col in row.index:
            col_str = str(col).lower().strip()
            
            # Skip empty or non-uniform columns
            if not col_str or col_str in ['', 'nan', 'none']:
                continue
                
            # Check if this looks like a prenda column
            if self._is_prenda_column(col_str):
                uniform_columns.append(col)
        
        return uniform_columns
    
    def _is_prenda_column(self, column_name: str) -> bool:
        """Check if a column name represents a prenda type.
        
        Supports both old format (OCCUPATION_PRENDA) and new format (LOCATION_OCCUPATION_PRENDA).
        """
        # Location group prefixes (new format)
        location_indicators = ['lima_ica', 'patios_comida', 'villa_steakhouse']
        
        # Occupation indicators (both old and new format)
        occupation_indicators = [
            'salon', 'delivery', 'packer', 'bar', 'seguridad', 'produccion', 
            'anfitrion', 'anfitrionaje', 'caja', 'mantenimiento', 'administracion',
            'auditoria', 'counter', 'corredor'
        ]
        
        # Prenda type indicators
        prenda_indicators = [
            'camisa', 'blusa', 'polo', 'casaca', 'gorra', 'mandilon', 'andarin',
            'pechera', 'saco', 'pantalon', 'garibaldi', 'chaqueta', 'gorro',
            'chaleco', 'saco_h', 'saco_m', 'polo_manga_corta'
        ]
        
        column_lower = column_name.lower()
        
        # Check if column has a location prefix (new format)
        has_location = any(loc in column_lower for loc in location_indicators)
        
        # Check for occupation or prenda indicator
        has_occupation = any(occ in column_lower for occ in occupation_indicators)
        has_prenda = any(prenda in column_lower for prenda in prenda_indicators)
        
        # New format: must have location prefix AND (occupation OR prenda)
        if has_location:
            return has_occupation or has_prenda
        
        # Old format: just needs occupation or prenda indicator
        return has_occupation or has_prenda
    
    def _normalize_prenda_type(self, column_name: str) -> str:
        """
        Normalize column name to standard prenda type.
        
        Handles both old format (OCCUPATION_PRENDA) and new format (LOCATION_OCCUPATION_PRENDA).
        Also handles compound prenda names like SACO_H, POLO_MANGA_CORTA.
        
        Examples:
            - "DELIVERY_POLO" -> "POLO"
            - "LIMA_ICA_DELIVERY_POLO" -> "POLO"  
            - "LIMA_ICA_ANFITRIONAJE_SACO_H" -> "SACO"
            - "PATIOS_COMIDA_PRODUCCION_POLO_MANGA_CORTA" -> "POLO"
        """
        col_upper = str(column_name).upper().strip()
        
        # List of known prenda types to extract
        known_prendas = [
            'CAMISA', 'BLUSA', 'POLO', 'CASACA', 'GORRA', 'MANDILON', 'ANDARIN',
            'PECHERA', 'SACO', 'PANTALON', 'GARIBALDI', 'CHAQUETA', 'GORRO', 'CHALECO',
            'CORBATA'  # Added CORBATA
        ]
        
        # Check if any known prenda type is in the column name
        for prenda in known_prendas:
            if prenda in col_upper:
                return prenda
        
        # Fallback: For new format, try splitting and looking for known parts
        if '_' in col_upper:
            parts = col_upper.split('_')
            # Search from the end for a known prenda
            for part in reversed(parts):
                if part in known_prendas:
                    return part
            # If no known prenda found, return the last meaningful part
            # Skip suffix modifiers like H, M, CORTA, numeric suffixes, etc.
            skip_suffixes = {'H', 'M', 'CORTA', 'MANGA', '1', '2', '3', '4', '5'}
            for part in reversed(parts):
                if part and part not in skip_suffixes and not part.isdigit():
                    return part
        
        # Fallback for legacy format - try mapping
        col_lower = col_upper.lower()
        mapping = {
            'camisa': 'CAMISA',
            'blusa': 'BLUSA', 
            'polo': 'POLO',
            'casaca': 'CASACA',
            'gorra': 'GORRA',
            'mandilon': 'MANDILON',
            'mandilón': 'MANDILON',
            'andarin': 'ANDARIN',
            'andarín': 'ANDARIN',
            'pechera': 'PECHERA',
            'saco': 'SACO',
            'pantalon': 'PANTALON',
            'pantalón': 'PANTALON',
            'garibaldi': 'GARIBALDI',
            'chaqueta': 'CHAQUETA',
            'gorro': 'GORRO'
        }
        
        for key, value in mapping.items():
            if key in col_lower:
                return value
        
        return col_upper.replace(' ', '_')
    
    def _get_display_name(self, prenda_type: str) -> str:
        """
        Get display name for prenda type.
        For document generation, we only show the prenda name (not the occupation prefix).
        Example: "POLO" -> "Polo", not "DELIVERY_POLO" -> "Delivery Polo"
        """
        # prenda_type is already just the prenda part (e.g., "POLO") after _normalize_prenda_type
        # Just title case it for display
        display = prenda_type.replace('_', ' ').title()
        return display
    
    def _format_prenda_string(self, display_name: str, talla_superior: str) -> str:
        """Format prenda string for display."""
        # Items that don't need "TALLA" prefix (typically one-size items)
        no_talla_items = {"ANDARIN", "MANDILON", "GORRA", "PECHERA", "PECHERA BAR", "GARIBALDI"}
        
        if display_name.upper() in no_talla_items:
            return display_name
        else:
            return f"{display_name} TALLA {talla_superior}"

    def _find_in_row(self, row: pd.Series, keys: List[str]) -> Optional[str]:
        """Find a value in a row by searching for column names that contain any of the specified keys."""
        lowered = {str(k).lower(): k for k in row.index}
        
        # First try exact matches
        for needle in keys:
            if needle.lower() in lowered:
                val = row[lowered[needle.lower()]]
                # Handle Series values (from duplicate column names)
                if isinstance(val, pd.Series):
                    val = val.iloc[0] if len(val) > 0 else None
                if pd.notna(val) and str(val).strip():
                    return str(val).strip()
        
        # Then try partial matches
        for key in lowered:
            for needle in keys:
                if needle.lower() in key:
                    val = row[lowered[key]]
                    # Handle Series values (from duplicate column names)
                    if isinstance(val, pd.Series):
                        val = val.iloc[0] if len(val) > 0 else None
                    if pd.notna(val) and str(val).strip():
                        return str(val).strip()
        
        return None

    def _extract_name(self, row: pd.Series) -> str:
        lowered = {str(k).lower(): k for k in row.index}
        combined = None
        for key in lowered:
            if "nombre" in key and "apellido" in key:
                combined = row[lowered[key]]
                # Handle Series values (from duplicate column names)
                if isinstance(combined, pd.Series):
                    combined = combined.iloc[0] if len(combined) > 0 else None
                break
        if pd.notna(combined) and str(combined).strip():
            return str(combined).strip()
        first = None
        last = None
        for key in lowered:
            if first is None and ("nombre" in key or "name" in key):
                first = row[lowered[key]]
                # Handle Series values (from duplicate column names)
                if isinstance(first, pd.Series):
                    first = first.iloc[0] if len(first) > 0 else None
            if last is None and ("apellido" in key or "last" in key):
                last = row[lowered[key]]
                # Handle Series values (from duplicate column names)
                if isinstance(last, pd.Series):
                    last = last.iloc[0] if len(last) > 0 else None
        name = ""
        if pd.notna(first):
            name = str(first).strip()
        if pd.notna(last):
            name = (name + " " + str(last).strip()).strip()
        return name

    def _get_uniform_columns_from_row(self, row: pd.Series) -> List[str]:
        """
        Get relevant uniform columns for this row based on occupation group.
        This prevents 'cross-talk' where an employee in one role gets items from another column group.
        """
        all_columns = [col for col in row.index if col not in ['apellidos_y_nombres', 'dni', 'cargo', 'fecha_ingrese', 
                                                             'talla_zapato', 'talla_pantalon', 'talla_prenda_superior']]
        
        # Get cargo from row
        cargo = self._find_in_row(row, ["cargo"]) or ""
        if not cargo:
            return all_columns

        # Normalize cargo
        from cargos.core.constants import OCCUPATION_GROUP_MAPPING, VILLA_STEAKHOUSE_COLUMN_MAPPING
        
        # Check explicit mapping first
        cargo_upper = str(cargo).upper().strip()
        
        # Find the group for this occupation
        # We try exact match, then partial match
        target_group = None
        
        # 1. Try exact match in mapping
        if cargo_upper in OCCUPATION_GROUP_MAPPING:
            target_group = OCCUPATION_GROUP_MAPPING[cargo_upper]
        
        # 2. Try partial match (e.g. "MOZO (EVENTUAL)" -> matches "MOZO")
        if not target_group:
            for key, group in OCCUPATION_GROUP_MAPPING.items():
                if key in cargo_upper:
                    target_group = group
                    break
        
        # If no group determined, return all columns (fallback)
        if not target_group:
            return all_columns
            
        # Filter columns that belong to this group
        filtered_columns = []
        for col in all_columns:
            # Column format is usually LOCATION_GROUP_ITEM (e.g. LIMA_ICA_SALON_CAMISA)
            # We check if the group name (e.g. "SALON") appears in the column name
            if f"_{target_group}_" in str(col).upper():
                filtered_columns.append(col)
        
        # Special Logic for Villa Steakhouse / San Isidro
        # If the location is Villa Steakhouse/San Isidro, they might use 'CORREDOR' items even if they are 'SALON'
        # But our OCCUPATION_GROUP_MAPPING says 'CORREDOR' -> 'CORREDOR'
        # So we should be careful. 
        # Actually, the user said: "if the tienda is SAN ISIDRO then you can assume that theyll use salon y corredor"
        # So if we are in San Isidro, and the person is SALON, we should ALSO include CORREDOR columns.
        
        # Check metadata from the worksheet this row belongs to?
        # That's hard because we just have the row here.
        # But we can infer from the column names if they are VILLA_STEAKHOUSE columns
        is_villa = any("VILLA_STEAKHOUSE" in str(c).upper() for c in all_columns)
        
        if is_villa and target_group == "SALON":
             # Add CORREDOR columns too
             for col in all_columns:
                 if "_CORREDOR_" in str(col).upper():
                     filtered_columns.append(col)

        if not filtered_columns:
             # If filtering removed everything (maybe mismatch in naming), fall back to all
             self.logger.warning(f"Filtering for group '{target_group}' resulted in 0 columns for '{cargo}'. Using all columns.")
             return all_columns
             
        return filtered_columns

    def _sanitize_name(self, s: str) -> str:
        return "".join(ch for ch in s if ch.isalnum() or ch in ("_", "-", " ")).strip().replace(" ", "_")

    def _file_stub(self, ctx: Dict[str, Any]) -> str:
        parts = [ctx.get("nombre", "").strip(), ctx.get("cargo", "").strip()]
        return self._sanitize_name("_".join([p for p in parts if p]))


# ConfigService removed - use ConfigManager directly from config_manager module
